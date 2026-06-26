"""文件解析服务 — 支持 docx / pdf / xlsx / txt"""
from pathlib import Path
from docx import Document
from PyPDF2 import PdfReader
from openpyxl import load_workbook


def parse_file(file_path: str) -> str:
    """
    解析上传文件，提取文本内容

    Args:
        file_path: 文件路径

    Returns:
        提取的文本内容
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".docx":
        return _parse_docx(path)
    elif suffix == ".pdf":
        return _parse_pdf(path)
    elif suffix in (".xlsx", ".xls"):
        return _parse_xlsx(path)
    elif suffix in (".txt", ".md"):
        return _parse_txt(path)
    else:
        raise ValueError(f"不支持的文件格式: {suffix}（支持 docx/pdf/xlsx/txt）")


def _parse_docx(path: Path) -> str:
    """解析 Word 文档"""
    doc = Document(str(path))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs)


def _parse_pdf(path: Path) -> str:
    """解析 PDF 文档"""
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text.strip())
    return "\n\n".join(pages)


def _parse_xlsx(path: Path) -> str:
    """解析 Excel 文档"""
    wb = load_workbook(str(path), data_only=True)
    sheets = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(cell) if cell is not None else "" for cell in row]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            sheets.append(f"## {sheet_name}\n" + "\n".join(rows))
    return "\n\n".join(sheets)


def _parse_txt(path: Path) -> str:
    """解析文本文件"""
    encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
    for enc in encodings:
        try:
            return path.read_text(encoding=enc)
        except (UnicodeDecodeError, IOError):
            continue
    raise ValueError(f"无法解码文件: {path.name}")
